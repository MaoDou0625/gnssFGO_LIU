# -*- coding: utf-8 -*-
"""Build the v2.3 core framework Word guide for thesis drafting."""

from __future__ import annotations

import subprocess
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
CFG_PATH = REPO_ROOT / "config" / "default_offline.cfg"
DOCX_PATH = REPO_ROOT / "docs" / "offline_lc_minimal_v2_3_core_framework.docx"


GROUP_ORDER = [
    "输入输出与运行控制",
    "时间线、同步与重力",
    "IMU噪声、bias与误差状态",
    "初始对准与静止约束",
    "GNSS/RTK质量控制",
    "Stage1航向精化",
    "Stage2速度与车辆约束",
    "Stage3高程参考与最终策略",
    "RTK中断、漂移与恢复",
    "垂向运动一致性",
    "Body-Z、NHC与路面状态",
    "垂向jump/bias处理",
    "优化器与通用末端参数",
    "其他兼容参数",
]


PARAMETER_NOTES = {
    "imu_path": "IMU输入文本路径。",
    "gnss_path": "GNSS/RTK解算结果输入路径。",
    "output_dir": "默认输出目录；正式运行会写入轨迹、摘要、诊断和配置快照。",
    "enable_gnss": "是否使用GNSS/RTK观测。",
    "enable_gp_interpolated_gnss": "是否使用GP插值形式对齐GNSS观测与状态节点。",
    "write_debug_csv": "写出GNSS残差、对齐等调试CSV。",
    "enable_stage_attitude_debug_export": "写出跨stage姿态参考调试轨迹。",
    "write_segment_error_diagnostics": "写出分段误差诊断。",
    "state_frequency_hz": "优化状态节点频率。",
    "error_state_frequency_hz": "误差状态输出频率。",
    "state_meas_sync_lower_bound_s": "状态与观测同步的下界。",
    "state_meas_sync_upper_bound_s": "状态与观测同步的上界。",
    "gravity_mps2": "重力常数。",
    "imu_sigma_acc_ug": "加速度计白噪声，配置中使用micro-g单位。",
    "imu_sigma_gyro_dph": "陀螺白噪声，配置中使用deg/hour单位。",
    "integration_sigma": "IMU预积分数值积分噪声。",
    "bias_acc_sigma_ug": "加速度计bias随机游走尺度。",
    "bias_gyro_sigma_dph": "陀螺bias随机游走尺度。",
    "bias_acc_prior_sigma_ug": "初始加速度计bias先验。",
    "bias_gyro_prior_sigma_dph": "初始陀螺bias先验。",
    "enable_global_acc_bias": "启用全局加速度计bias变量。",
    "enable_global_gyro_bias": "启用全局陀螺bias变量。",
    "enable_vertical_acc_bias_gm_process": "启用垂向加速度bias Gauss-Markov过程。",
    "vertical_acc_bias_tau_s": "垂向加速度bias GM相关时间。",
    "vertical_acc_bias_sigma_ug": "垂向加速度bias GM强度。",
    "enable_stage1_yaw_refinement": "启用Stage1初始yaw迭代精化。",
    "stage1_yaw_refinement_max_iterations": "Stage1 yaw精化最大迭代次数。",
    "stage1_heading_window_s": "RTK heading估计窗口。",
    "stage1_heading_min_displacement_m": "heading估计最小水平位移。",
    "stage1_yaw_update_max_rad": "单轮yaw更新最大幅度。",
    "enable_stage1_outage_body_y_envelope": "启用RTK中断段横向速度包络约束。",
    "enable_stage2_velocity_optimization": "启用Stage2速度优化。",
    "enable_stage2_vehicle_nhc_constraint": "Stage2中启用车辆坐标系NHC。",
    "stage2_attitude_hold_sigma_rad": "Stage3继承Stage2姿态时使用的hold sigma。",
    "stage2_horizontal_position_hold_sigma_m": "Stage3继承Stage2水平位置的hold sigma。",
    "stage2_horizontal_velocity_hold_sigma_mps": "Stage3继承Stage2水平速度的hold sigma。",
    "stage2_vehicle_y_nhc_velocity_sigma_mps": "Stage2车辆y向速度NHC sigma。",
    "stage2_vehicle_y_nhc_displacement_sigma_m": "Stage2车辆y向位移NHC sigma。",
    "enable_stage2_lowfreq_vertical_reference_optimization": "旧Stage2低频高程参考流程；当前默认关闭。",
    "enable_stage3_vertical_reference_optimization": "启用Stage3高程参考优化主流程。",
    "stage3_vertical_reference_smoothing_method": "Stage3内嵌参考的平滑方法。",
    "stage3_vertical_reference_lowpass_cutoff_hz": "Stage3低通fallback截止频率。",
    "stage3_vertical_reference_spline_knot_spacing_m": "Stage3 spline baseline knot距离间隔。",
    "stage3_vertical_reference_spline_smooth_lambda": "Stage3 spline平滑项权重。",
    "stage3_vertical_reference_spline_anchor_weight": "Stage3 spline anchor权重。",
    "stage3_vertical_reference_spline_slope_weight": "Stage3 spline斜率平滑权重。",
    "stage3_vertical_reference_constraint_mode": "配置文件基础值；final pass会被策略覆盖为envelope。",
    "stage3_vertical_anchor_sigma_m": "Gaussian路径和共享参考缺省sigma的兼容参数。",
    "stage3_vertical_envelope_half_width_m": "Stage3高程包络半宽；final pass强制为5 mm。",
    "stage3_vertical_envelope_sigma_m": "Stage3 envelope factor sigma；final pass强制为3 mm。",
    "enable_stage3_vertical_envelope_center_pull": "Stage3 center-pull；final pass强制关闭。",
    "enable_stage3_stage2_vertical_increment_hold": "约束Stage3继承Stage2相邻垂向增量。",
    "stage3_stage2_vertical_increment_sigma_m": "非jump区Stage2垂向增量继承sigma。",
    "stage3_stage2_vertical_increment_jump_sigma_m": "jump区Stage2垂向增量继承sigma。",
    "enable_stage3_stage2_jump_shape_hold": "约束Stage3继承Stage2 jump窗口内相对高程形状。",
    "stage3_stage2_jump_shape_sigma_m": "Stage2 jump形状继承sigma。",
    "enable_stage3_jump_velocity_smoothness_regularizer": "旧Stage3 jump速度平滑正则；当前final关闭。",
    "enable_stage3_jump_height_highfreq_deadband": "旧Stage3 jump高频高程deadband；当前final关闭。",
    "enable_stage3_jump_adaptive_context_envelope": "旧Stage3 jump上下文自适应包络；当前final关闭。",
    "gnss_vertical_sigma_mode": "GNSS垂向sigma来源。",
    "gnss_vertical_fixed_sigma_m": "固定垂向GNSS sigma。",
    "vertical_constraint_mode": "普通GNSS垂向约束模式。",
    "gnss_vertical_reference_source": "垂向参考来源；Stage3 final会回到raw_rtk但raw GNSS因子被关闭。",
    "enable_rtk_vertical_drift_reference": "Stage1/2中的RTK垂向漂移参考；Stage3 final关闭。",
    "rtk_vertical_drift_correlation_time_s": "RTK垂向漂移相关时间。",
    "rtk_vertical_drift_sigma_m": "RTK垂向漂移sigma。",
    "rtk_vertical_white_noise_sigma_m": "RTK垂向白噪声sigma。",
    "enable_rtk_outage_segmented_batch": "启用RTK中断分段batch优化。",
    "stage3_disable_rtk_outage_segmented_batch": "Stage3包装开关；final策略关闭内层RTK outage递归。",
    "rtk_outage_segmented_batch_max_outages": "默认最多处理的RTK outage个数。",
    "enable_rtk_outage_boundary_constraints": "启用outage边界位置、速度、bias约束。",
    "enable_rtk_outage_baz_reestimate": "启用outage内ba_z重估规划。",
    "enable_rtk_outage_attitude_hold": "启用RTK outage姿态handoff/hold。",
    "enable_rtk_outage_velocity_delta_3d": "启用outage 3D速度delta；Stage3 final关闭。",
    "enable_vertical_velocity_delta_constraint": "启用垂向速度delta约束。",
    "vertical_velocity_delta_acc_sigma_ug": "垂向速度delta由加速度不确定度推导的sigma。",
    "vertical_velocity_delta_sigma_scale": "垂向速度delta整体sigma缩放。",
    "enable_vertical_motion_adaptive_reweighting": "启用垂向运动自适应重加权。",
    "enable_vertical_position_velocity_consistency_all_states": "启用全状态垂向位置-速度一致性约束。",
    "enable_vertical_position_velocity_window_consistency": "启用窗口级垂向位置-速度一致性约束。",
    "enable_attitude_reference_constraint": "启用姿态参考约束；Stage3 final关闭竞争项。",
    "enable_base_graph_tilt_reference_constraint": "启用base graph tilt参考；Stage3 final关闭。",
    "enable_body_z_nhc_constraint": "启用body-z NHC；Stage3 final关闭竞争项。",
    "enable_body_z_nhc_horizontal_leakage_correction": "启用body-z水平泄漏估计；Stage3 final关闭竞争项。",
    "enable_body_z_jump_detection": "启用body-z jump检测。",
    "enable_road_noise_state_baz_reestimate": "启用路面噪声状态触发的ba_z重估。",
    "enable_vertical_jump_bias": "启用jump bias约束；当前默认保留。",
    "enable_vertical_jump_impulse": "jump impulse模型；当前默认关闭。",
    "enable_vertical_jump_masked_imu": "jump区IMU mask；当前默认关闭。",
    "enable_vertical_jump_segmented_bias": "jump分段bias；当前默认关闭。",
    "gnss_position_noise_model": "GNSS位置因子的鲁棒噪声模型。",
    "gnss_position_robust_param": "GNSS鲁棒核参数。",
    "drop_non_rtkfix": "默认丢弃非RTKFIX普通GNSS样本。",
    "drop_no_solution": "丢弃NO_SOLUTION样本。",
    "drop_nonfinite_sigma": "丢弃sigma非有限样本。",
    "gnss_consistency_gate_mode": "GNSS一致性门限模式。",
    "lm_lambda_initial": "Levenberg-Marquardt初始lambda。",
    "lm_max_iterations": "Levenberg-Marquardt最大迭代次数。",
}


MODULE_ROWS = [
    ("入口/调度", "offline_lc_runner", "src/offline_runner/main.cpp", "读取cfg与覆盖项，调用OfflineBatchRunner并写出结果。", "默认主入口"),
    ("入口/共享高程", "offline_lc_shared_vertical_reference_builder", "src/shared_vertical_reference_builder/main.cpp", "从多成员Stage2轨迹和RTKFIX生成统一距离域z_shared(s)。", "手动共享流程"),
    ("入口/Stage3-only", "offline_lc_stage3_runner", "src/stage3_runner/main.cpp", "读取Stage2 trajectory和共享高程参考，只执行Stage3 final。", "手动共享流程"),
    ("配置", "Config", "src/common/Config.cpp; include/.../Config.h", "解析default_offline.cfg、单位迁移、校验和config_snapshot写出。", "启用"),
    ("数据IO", "TextDatasetLoader", "src/io/TextDatasetLoader.cpp", "加载IMU/GNSS文本数据并形成DataSet。", "启用"),
    ("数据IO", "TrajectoryCsvReader", "src/io/TrajectoryCsvReader.cpp", "读取Stage2 trajectory作为后续reference。", "Stage3-only启用"),
    ("输出", "ResultWriter", "src/common/ResultWriter.cpp", "写trajectory、summary、diagnostics、config_snapshot。", "启用"),
    ("时间线", "GraphTimelineBuilder", "src/core/GraphTimelineBuilder.cpp", "建立统一状态节点时间线和动态起点。", "启用"),
    ("初始化", "TrajectoryInitializer", "src/core/TrajectoryInitializer.cpp", "构造位置、速度、姿态、bias初值。", "启用"),
    ("初始化", "StaticImuAlignment", "src/core/StaticImuAlignment.cpp", "初始静止段IMU对准。", "条件启用"),
    ("IMU", "ImuIntegrationUtils", "src/core/ImuIntegrationUtils.cpp", "IMU预积分、相对传播与辅助计算。", "启用"),
    ("GNSS", "GnssFactorBuilder", "src/core/GnssFactorBuilder.cpp", "构建GNSS位置/垂向/速度因子；Stage3 reference存在时关闭raw GNSS因子。", "Stage1/2启用，Stage3 raw关闭"),
    ("GNSS", "GnssPreOutageQualityOverride", "src/core/GnssPreOutageQualityOverride.cpp", "outage前GNSS质量覆盖。", "启用"),
    ("GNSS", "GnssVerticalReferenceSelector", "src/core/GnssVerticalReferenceSelector.cpp", "选择垂向GNSS参考。", "条件启用"),
    ("Stage1", "Stage1YawRefinementRunner", "src/core/Stage1YawRefinementRunner.cpp", "迭代修正初始yaw并输出可传递的Stage1参考。", "启用"),
    ("Stage1", "Stage1YawBranchResolver", "src/core/Stage1YawBranchResolver.cpp", "处理90/180度分支与未收敛情形。", "条件启用"),
    ("Stage1", "Stage1SourceReferencePolicy", "src/core/Stage1SourceReferencePolicy.cpp", "判断Stage1 source是否可安全传给Stage2。", "Stage2启用"),
    ("Stage1", "Stage1OutageLateralVelocityEnvelopeEstimator", "src/core/Stage1OutageLateralVelocityEnvelopeEstimator.cpp", "估计outage内车辆横向速度包络。", "启用"),
    ("Stage1", "Stage1OutageBodyYEnvelopeConstraintBuilder", "src/core/Stage1OutageBodyYEnvelopeConstraintBuilder.cpp", "把body-y envelope加入图优化。", "条件启用"),
    ("Stage2", "Stage2VelocityOptimizationRunner", "src/core/Stage2VelocityOptimizationRunner.cpp", "以Stage1为source，生成Stage2VelocityReference并运行Stage2优化。", "启用"),
    ("Stage2", "Stage2VelocityReference", "src/core/Stage2VelocityReference.cpp", "保存rotation-native姿态与状态参考，供Stage3继承。", "启用"),
    ("Stage2", "Stage2VehicleNHCConstraintBuilder", "src/core/Stage2VehicleNHCConstraintBuilder.cpp", "车辆y向速度/位移NHC。", "Stage2启用，Stage3关闭"),
    ("Stage2", "Stage2AttitudeHoldBuilder", "src/core/Stage2AttitudeHoldBuilder.cpp", "Stage3中保持Stage2姿态。", "Stage3启用"),
    ("Stage2", "Stage2HorizontalHoldBuilder", "src/core/Stage2HorizontalHoldBuilder.cpp", "Stage3中保持Stage2水平位置和水平速度。", "Stage3启用"),
    ("Stage3", "Stage3VerticalReferenceOptimizationRunner", "src/core/Stage3VerticalReferenceOptimizationRunner.cpp", "内嵌Stage3主流程：Stage2 source、规划高程参考、final pass。", "启用"),
    ("Stage3", "Stage3HeightOptimizationPolicy", "src/core/Stage3HeightOptimizationPolicy.cpp", "生成Stage3 final子配置并强制低频高程策略。", "启用"),
    ("Stage3", "Stage3VerticalReferenceProfilePlanner", "src/core/Stage3VerticalReferenceProfilePlanner.cpp", "从Stage2轨迹规划per-state高程参考。", "内嵌Stage3启用"),
    ("Stage3", "Stage3VerticalReferenceSmoother", "src/core/Stage3VerticalReferenceSmoother.cpp", "低通或spline baseline高程平滑。", "启用"),
    ("Stage3", "Stage3VerticalReferenceConstraintBuilder", "src/core/Stage3VerticalReferenceConstraintBuilder.cpp", "构建Stage3高程reference envelope/Gaussian因子。", "Stage3启用"),
    ("Stage3", "Stage3Stage2IncrementHoldConstraintBuilder", "src/core/Stage3Stage2IncrementHoldConstraintBuilder.cpp", "约束Stage3继承Stage2相邻垂向增量。", "Stage3启用"),
    ("Stage3", "Stage3Stage2JumpShapeHoldConstraintBuilder", "src/core/Stage3Stage2JumpShapeHoldConstraintBuilder.cpp", "约束Stage3继承Stage2 jump相对形状。", "Stage3启用"),
    ("共享高程", "SharedVerticalReferenceBuilder", "src/core/SharedVerticalReferenceBuilder.cpp", "构造统一参考线和z_shared(s)。", "手动共享流程"),
    ("共享高程", "Stage3SharedReferenceMapper", "src/core/Stage3SharedReferenceMapper.cpp", "把Stage2状态投影到共享参考线并生成Stage3 reference。", "Stage3-only启用"),
    ("共享高程", "Stage3SharedReferenceDeltaSmoother", "src/core/Stage3SharedReferenceDeltaSmoother.cpp", "平滑z_shared与Stage2之间的delta。", "Stage3-only启用"),
    ("RTK outage", "RtkOutageWindowPlanner", "src/core/RtkOutageWindowPlanner.cpp", "检测和规划RTK outage窗口。", "Stage1/2启用"),
    ("RTK outage", "RtkOutageSegmentedBatchRunner", "src/core/RtkOutageSegmentedBatchRunner.cpp", "pre/outage/post分段batch优化。", "Stage1/2启用，Stage3关闭递归"),
    ("RTK outage", "RtkOutageBoundaryConstraintBuilder", "src/core/RtkOutageBoundaryConstraintBuilder.cpp", "构造outage边界位置/速度/bias约束。", "启用"),
    ("RTK outage", "RtkOutageRecoveryConstraintBuilder", "src/core/RtkOutageRecoveryConstraintBuilder.cpp", "构造RTK恢复段位置/速度/姿态恢复约束。", "条件启用"),
    ("RTK drift", "RtkVerticalDriftReferenceEstimator", "src/core/RtkVerticalDriftReferenceEstimator.cpp", "估计RTK垂向漂移参考。", "Stage1/2启用，Stage3关闭"),
    ("Body-Z/NHC", "BodyZWindowPipeline", "src/core/BodyZWindowPipeline.cpp", "body-z窗口生成、jump/seed诊断与NHC前处理。", "条件启用"),
    ("Body-Z/NHC", "BodyZBidirectionalJumpDetector", "src/core/BodyZBidirectionalJumpDetector.cpp", "双向jump检测。", "启用"),
    ("Body-Z/NHC", "RoadNoiseStateEstimator", "src/core/RoadNoiseStateEstimator.cpp", "路面噪声状态分段。", "启用"),
    ("Body-Z/NHC", "BodyZBiasReestimatePlanner", "src/core/BodyZBiasReestimatePlanner.cpp", "规划body-z bias重估窗口。", "条件启用"),
    ("垂向约束", "VerticalMotionConstraintBuilder", "src/core/VerticalMotionConstraintBuilder.cpp", "构建垂向速度delta、运动一致性和相关约束。", "启用"),
    ("垂向约束", "VerticalAdaptiveReweightingLoop", "src/core/VerticalAdaptiveReweightingLoop.cpp", "垂向运动自适应重加权。", "启用"),
    ("垂向约束", "VerticalPositionVelocityConsistencyConstraintBuilder", "src/core/VerticalPositionVelocityConsistencyConstraintBuilder.cpp", "垂向位置-速度一致性约束。", "启用"),
    ("垂向jump", "VerticalJumpBiasConstraintBuilder", "src/core/VerticalJumpBiasConstraintBuilder.cpp", "jump bias因子。", "启用"),
    ("垂向jump", "VerticalJumpImpulseConstraintBuilder", "src/core/VerticalJumpImpulseConstraintBuilder.cpp", "jump impulse因子。", "默认关闭"),
    ("诊断", "RunDiagnosticsBuilder", "src/core/RunDiagnosticsBuilder.cpp", "汇总运行诊断和summary字段。", "启用"),
    ("诊断", "ResidualContributionAnalyzer", "src/core/ResidualContributionAnalyzer.cpp", "分析残差模块/因子贡献。", "启用"),
]


STAGE3_FINAL_POLICY = [
    ("stage3_vertical_reference_constraint_mode", "gaussian", "envelope", "final pass使用5 mm包络门限，不再直接高斯拉向中心。"),
    ("stage3_vertical_anchor_sigma_m", "0.001", "0.001", "保留为兼容Gaussian路径和共享参考缺省sigma。"),
    ("stage3_vertical_envelope_half_width_m", "0.008", "0.005", "包络半宽收窄到5 mm。"),
    ("stage3_vertical_envelope_sigma_m", "0.003", "0.003", "包络外残差sigma为3 mm。"),
    ("enable_stage3_vertical_envelope_center_pull", "true", "false", "关闭center-pull，避免引入短波高频毛刺。"),
    ("enable_stage3_stage2_vertical_increment_hold", "true", "true", "继承Stage2相邻垂向增量。"),
    ("stage3_stage2_vertical_increment_sigma_m", "0.0002", "0.0002", "非jump区增量继承强约束。"),
    ("stage3_stage2_vertical_increment_jump_sigma_m", "0.0005", "0.0005", "jump区增量继承稍放松。"),
    ("enable_stage3_stage2_jump_shape_hold", "true", "true", "继承Stage2 jump相对形状。"),
    ("stage3_stage2_jump_shape_sigma_m", "0.0005", "0.0005", "jump形状继承sigma。"),
    ("enable_stage3_jump_velocity_smoothness_regularizer", "false", "false", "旧jump速度平滑正则保持关闭。"),
    ("enable_stage3_jump_height_highfreq_deadband", "false", "false", "旧高频高程deadband保持关闭。"),
    ("enable_stage3_jump_adaptive_context_envelope", "false", "false", "旧上下文自适应包络保持关闭。"),
    ("enable_rtk_vertical_drift_reference", "true", "false", "Stage3 final不再递归使用RTK drift垂向参考。"),
    ("enable_rtk_outage_smoothing", "true", "false", "Stage3 final关闭RTK outage垂向平滑竞争项。"),
    ("enable_late_static_detection", "true", "false", "Stage3 final关闭末端静止raw RTK拉拽。"),
    ("enable_initial_static_rtk_height_reference", "true", "false", "Stage3 final关闭初始静止RTK高程参考。"),
    ("enable_stage2_vehicle_nhc_constraint", "true", "false", "Stage3 final关闭车辆NHC竞争项。"),
    ("enable_body_z_nhc_constraint", "true", "false", "Stage3 final关闭body-z NHC竞争项。"),
    ("enable_attitude_reference_constraint", "true", "false", "Stage3 final关闭姿态参考竞争项，转由Stage2 hold保持。"),
]


OUTPUT_ROWS = [
    ("trajectory.csv", "最终轨迹；含时间、位置、速度、姿态输出、bias等。"),
    ("summary.txt", "运行摘要；包含Stage1/2/3关键统计。"),
    ("data_summary.txt", "输入数据统计。"),
    ("config_snapshot.cfg", "本次运行的实际配置快照，是复现实验的关键文件。"),
    ("gnss_residuals.csv", "GNSS残差调试输出。"),
    ("stage1_yaw_refinement_diagnostics.csv", "Stage1 yaw迭代诊断。"),
    ("stage1_yaw_residual_module_contributions.csv", "Stage1残差模块贡献。"),
    ("stage_attitude_debug_trajectory.csv", "跨stage姿态调试轨迹。"),
    ("rtk_outage_windows.csv", "RTK outage窗口。"),
    ("rtk_outage_batch_segments.csv", "RTK outage分段batch信息。"),
    ("body_z_seed_jump_windows.csv", "body-z seed jump窗口。"),
    ("body_z_bias_reestimate_segments.csv", "body-z/road-state bias重估窗口。"),
    ("road_noise_state_segments.csv", "路面噪声状态分段。"),
    ("stage3_vertical_reference_diagnostics.csv", "Stage3高程参考残差、包络和skip原因。"),
    ("stage3_stage2_increment_hold_diagnostics.csv", "Stage3继承Stage2垂向增量诊断。"),
    ("stage3_stage2_jump_shape_hold_diagnostics.csv", "Stage3继承Stage2 jump形状诊断。"),
    ("vertical_velocity_delta_diagnostics.csv", "垂向速度delta约束诊断。"),
    ("vertical_jump_bias_diagnostics.csv", "垂向jump bias诊断。"),
    ("shared_vertical_reference.csv", "共享距离域z_shared(s)。"),
    ("shared_reference_line.csv", "共享投影参考线。"),
    ("shared_vertical_reference_projection_diagnostics.csv", "共享参考构造中的RTK/Stage2投影诊断。"),
]


def run_git(args: list[str]) -> str:
    try:
        return subprocess.check_output(
            ["git", *args],
            cwd=REPO_ROOT,
            text=True,
            encoding="utf-8",
            errors="replace",
        ).strip()
    except Exception:
        return "unknown"


def parse_config(path: Path) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, value = line.split("=", 1)
        rows.append((key.strip(), value.strip()))
    return rows


def group_for_key(key: str) -> str:
    if key in {"imu_path", "gnss_path", "output_dir"} or key.startswith("write_") or key in {
        "verbose",
        "enable_gnss",
        "enable_gp_interpolated_gnss",
    }:
        return "输入输出与运行控制"
    if key.startswith("state_") or key.startswith("processing_") or key in {
        "gravity_mps2",
        "gnss_time_offset_s",
    }:
        return "时间线、同步与重力"
    if (
        key.startswith("imu_")
        or "bias" in key and not key.startswith("vertical_jump")
        or key.startswith("error_state")
        or key.startswith("segment_feedback")
        or key.startswith("tau_")
        or key.startswith("global_")
        or key.startswith("enable_global")
        or key.startswith("vertical_acc_bias")
        or key == "integration_sigma"
        or key.startswith("error_process")
        or key.startswith("bias_process")
    ):
        return "IMU噪声、bias与误差状态"
    if key.startswith("initial_static") or key.startswith("initial_dynamic") or key.startswith("late_static") or key in {
        "stationary_window_s",
        "stationary_acc_tolerance_mps2",
        "stationary_gyro_threshold_radps",
        "static_alignment_duration_s",
        "imu_dual_vector_window_s",
        "imu_dual_vector_min_sample_count",
        "imu_dual_vector_min_cross_norm",
        "prefer_imu_initial_yaw",
        "enable_initial_yaw_override",
        "initial_yaw_override_rad",
        "yaw_min_distance_m",
        "fallback_initial_yaw_rad",
    }:
        return "初始对准与静止约束"
    if key.startswith("stage1_") or key == "enable_stage1_yaw_refinement":
        return "Stage1航向精化"
    if key.startswith("stage2_") or key == "enable_stage2_velocity_optimization" or key == "enable_stage2_vehicle_nhc_constraint":
        return "Stage2速度与车辆约束"
    if key.startswith("stage3_") or key.startswith("enable_stage3_"):
        return "Stage3高程参考与最终策略"
    if key.startswith("rtk_outage") or key.startswith("enable_rtk_outage") or key.startswith("rtk_vertical") or key.startswith("enable_rtk_vertical"):
        return "RTK中断、漂移与恢复"
    if key.startswith("vertical_velocity") or key.startswith("vertical_motion") or key.startswith("vertical_position") or key.startswith("vertical_constraint") or key.startswith("vertical_envelope") or key.startswith("enable_vertical_velocity") or key.startswith("enable_vertical_motion") or key.startswith("enable_vertical_position") or key.startswith("enable_vertical_envelope"):
        return "垂向运动一致性"
    if key.startswith("body_z") or key.startswith("enable_body_z") or key.startswith("road_noise") or key.startswith("enable_road_noise"):
        return "Body-Z、NHC与路面状态"
    if key.startswith("vertical_jump") or key.startswith("enable_vertical_jump"):
        return "垂向jump/bias处理"
    if key.startswith("gnss_") or key.startswith("enable_gnss") or key.startswith("drop_") or key in {
        "position_sigma_floor_m",
        "position_sigma_floor_horizontal_m",
        "position_sigma_floor_up_m",
        "position_sigma_ceiling_m",
        "rtkfix_scale",
        "rtkfloat_scale",
        "single_scale",
        "required_best_sol_status_code",
        "early_gnss_relaxation_duration_s",
        "early_gnss_relaxation_scale",
    }:
        return "GNSS/RTK质量控制"
    if key.startswith("lm_") or key.startswith("initial_position") or key.startswith("initial_roll") or key.startswith("initial_yaw_sigma") or key.startswith("initial_velocity"):
        return "优化器与通用末端参数"
    return "其他兼容参数"


def bool_status(value: str) -> str:
    lower = value.lower()
    if lower == "true":
        return "启用/true"
    if lower == "false":
        return "关闭/false"
    return "参数值"


def fit_text(text: str, max_len: int = 120) -> str:
    if len(text) <= max_len:
        return text
    return text[: max_len - 1] + "…"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = None
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None, size: int = 8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, size=size)
        set_cell_shading(hdr[idx], "D9EAF7")
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=size)
            if widths:
                cells[idx].width = Inches(widths[idx])
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def build_document() -> None:
    config_rows = parse_config(CFG_PATH)
    grouped: dict[str, list[tuple[str, str]]] = defaultdict(list)
    for key, value in config_rows:
        grouped[group_for_key(key)].append((key, value))

    git_describe = run_git(["describe", "--tags", "--always", "--dirty"])
    git_head = run_git(["rev-parse", "--short", "HEAD"])
    project_version = "2.3.0"

    doc = Document()
    style_document(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.text = "offline_lc_minimal v2.3 core framework"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("offline_lc_minimal v2.3\n核心算法框架与默认配置说明书")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"用于论文框架梳理 | project version {project_version} | {git_describe} | HEAD {git_head}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.name = "Microsoft YaHei"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_para(
        doc,
        "本文档基于当前仓库的v2.3工作线生成，核心依据为config/default_offline.cfg、"
        "Config.h/Config.cpp、OfflineBatchRunner以及Stage1/Stage2/Stage3相关策略模块。"
        "它不是运行结果报告，而是便于论文写作时抽取系统框架、模块职责和默认参数的技术说明。"
    )

    add_heading(doc, "1. 文档边界与版本依据", 1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("代码版本", f"project(... VERSION {project_version}); git describe: {git_describe}; HEAD: {git_head}"),
            ("正式默认配置", "config/default_offline.cfg。发布、复现实验和论文参数整理以此文件为准。"),
            ("代码兜底默认", "include/offline_lc_minimal/common/Config.h。它服务测试和临时构造路径，不是release cfg完整副本。"),
            ("配置解析/快照", "src/common/Config.cpp负责解析、校验、单位转换和config_snapshot.cfg写出。"),
            ("Stage3最终策略", "src/core/Stage3HeightOptimizationPolicy.cpp会在final pass前再次覆盖部分配置。"),
            ("共享高程流程", "Stage2-only -> shared z reference -> Stage3-only，使用两个专用CLI工具。"),
        ],
        widths=[1.7, 5.7],
    )
    add_para(
        doc,
        "论文中引用默认参数时，应区分三层含义：配置文件中的基础值、stage policy生成的子配置、"
        "单次实验输出目录中的config_snapshot.cfg。尤其是Stage3，高程final pass的权威值来自"
        "MakeStage3HeightOptimizationConfig()，不能仅抄default_offline.cfg里的基础值。"
    )

    add_heading(doc, "2. 核心算法框架", 1)
    add_para(
        doc,
        "系统采用离线因子图优化框架。状态节点按固定频率建立，IMU预积分负责相邻状态传播，GNSS/RTK、"
        "静止、车辆非完整约束、垂向运动一致性、jump/bias和stage间参考继承以因子形式加入图。"
        "优化目标可概括为最小化所有残差的加权平方和：min sum_i ||r_i(x)||^2_{W_i}。"
    )
    add_bullets(
        doc,
        [
            "状态主线：位置、速度、姿态、IMU bias以及垂向bias相关扩展变量。",
            "姿态内部传递：跨stage姿态参考使用rotation-native/SO(3)表示，yaw/pitch/roll主要作为输出和诊断字段。",
            "两层图结构：先优化base graph得到稳定参考，再叠加GNSS、垂向、outage、NHC和stage reference等约束形成最终图。",
            "递归stage调度：配置打开Stage3时，外层先生成Stage2 source，再生成高程参考，最后执行Stage3 final pass。",
            "当前v2.3核心原则：Stage2给出姿态、水平位置、水平速度和bias连续性；Stage3只做低频绝对高程校正。"
        ],
    )

    add_heading(doc, "3. 默认运行主流程", 1)
    add_numbered(
        doc,
        [
            "offline_lc_runner读取DefaultConfig、default_offline.cfg和命令行override，完成ValidateConfig。",
            "TextDatasetLoader加载IMU与GNSS，GNSS时间偏移、状态同步窗口和质量门限在此后参与样本筛选。",
            "OfflineBatchRunner检查stage开关。默认enable_stage3_vertical_reference_optimization=true，因此进入Stage3VerticalReferenceOptimizationRunner。",
            "Stage3 source run先通过MakeStage3HeightReferenceSourceConfig()关闭Stage3递归，只运行到Stage2并保留Stage2VelocityReference。",
            "Stage2VelocityOptimizationRunner内部调用Stage1YawRefinementRunner，必要时估计outage body-y envelope，再运行Stage2速度/NHC/垂向约束。",
            "Stage3VerticalReferenceProfilePlanner用Stage2轨迹构造低频高程参考；共享流程则由Stage3SharedReferenceMapper把z_shared(s)映射到本成员状态。",
            "MakeStage3HeightOptimizationConfig()生成final配置：关闭竞争性RTK垂向参考、姿态/NHC/水平重优化项，启用5 mm envelope和Stage2增量/形状继承。",
            "OfflineBatchRunner构建最终因子图并用Levenberg-Marquardt优化，ResultWriter写出trajectory、summary、diagnostics和config_snapshot。"
        ],
    )

    add_heading(doc, "4. Stage职责划分", 1)
    add_table(
        doc,
        ["阶段", "主要目标", "默认关键模块", "输出/传递内容"],
        [
            (
                "Stage1",
                "获得稳定初值和yaw参考，避免未收敛姿态传递。",
                "Stage1YawRefinementRunner, RtkHeadingAlignmentEstimator, Stage1YawBranchResolver",
                "Stage1 trajectory, reference states, yaw diagnostics, residual contributions",
            ),
            (
                "Stage2",
                "基于Stage1稳定姿态和RTK/GNSS信息优化速度、NHC和垂向运动连续性。",
                "Stage2VelocityOptimizationRunner, Stage2VehicleNHCConstraintBuilder, RtkOutageSegmentedBatchRunner",
                "Stage2VelocityReference, trajectory.csv, body_z_bias_reestimate_segments.csv",
            ),
            (
                "Shared z reference",
                "多成员共享同一距离域低频绝对高程目标。",
                "SharedVerticalReferenceBuilder, Stage3SharedReferenceMapper",
                "shared_vertical_reference.csv, shared_reference_line.csv, projection diagnostics",
            ),
            (
                "Stage3 final",
                "继承Stage2姿态/水平/速度/bias，只修正低频高程。",
                "Stage3HeightOptimizationPolicy, Stage3VerticalReferenceConstraintBuilder, Stage3Stage2IncrementHoldConstraintBuilder",
                "Stage3 trajectory, stage3_vertical_reference_diagnostics, increment/jump-shape diagnostics",
            ),
        ],
        widths=[0.9, 2.0, 2.5, 2.2],
    )

    add_heading(doc, "5. 默认启用模块清单", 1)
    add_para(
        doc,
        "下表按论文可描述的功能模块整理。状态列表示当前默认流程中的实际角色；某些模块虽然编译进库，"
        "但只在Stage1/2、Stage3-only或条件分支中启用。"
    )
    add_table(doc, ["类别", "模块", "代码位置", "职责", "默认状态"], MODULE_ROWS, widths=[0.9, 1.7, 2.2, 2.9, 1.1], size=7)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. Stage3 final策略覆盖", 1)
    add_para(
        doc,
        "default_offline.cfg中的Stage3字段描述了基础配置和参考规划默认值。真正进入Stage3 final优化前，"
        "MakeStage3HeightOptimizationConfig()会应用下列覆盖，确保Stage3只做低频高程校正。"
    )
    add_table(
        doc,
        ["参数", "cfg基础值", "Stage3 final值", "论文解释"],
        STAGE3_FINAL_POLICY,
        widths=[2.3, 1.2, 1.2, 3.0],
        size=7,
    )

    add_heading(doc, "7. 关键默认参数速查", 1)
    key_groups = [
        "输入输出与运行控制",
        "IMU噪声、bias与误差状态",
        "GNSS/RTK质量控制",
        "Stage1航向精化",
        "Stage2速度与车辆约束",
        "Stage3高程参考与最终策略",
        "RTK中断、漂移与恢复",
        "垂向运动一致性",
        "Body-Z、NHC与路面状态",
        "垂向jump/bias处理",
    ]
    for group in key_groups:
        rows = []
        for key, value in grouped.get(group, []):
            if key in PARAMETER_NOTES or key.startswith("enable_") or key.startswith("stage3_") or key.startswith("stage2_") or key.startswith("stage1_"):
                rows.append((key, value, bool_status(value), PARAMETER_NOTES.get(key, "该参数属于本组默认配置；完整值见附录参数清单。")))
        if not rows:
            continue
        add_heading(doc, group, 2)
        add_table(doc, ["参数", "默认值", "状态", "说明"], rows, widths=[2.2, 1.5, 1.0, 3.1], size=7)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "8. 输出与论文复现实验材料", 1)
    add_para(
        doc,
        "论文实验复现不应只保存trajectory.csv。每次正式运行都应保存config_snapshot.cfg、summary.txt、"
        "关键诊断CSV和共享参考文件；其中config_snapshot.cfg是确认stage policy和命令行override后实际参数的主要证据。"
    )
    add_table(doc, ["文件", "用途"], OUTPUT_ROWS, widths=[2.9, 4.7], size=8)

    add_heading(doc, "9. 可直接用于论文的框架表述", 1)
    add_bullets(
        doc,
        [
            "本文方法将GNSS/IMU离线组合定位建模为多约束因子图优化问题，通过IMU预积分建立连续运动先验，通过GNSS/RTK、车辆运动学、静止段和垂向一致性约束约束状态。",
            "为降低姿态误差向后续阶段传播，系统首先进行基于RTK航向的一阶段yaw精化，并对未收敛或双解情形进行分支判别。",
            "第二阶段以Stage1为姿态和状态参考，重点增强速度、车辆非完整约束、RTK中断分段和垂向bias连续性，形成后续高程修正的稳定基线。",
            "第三阶段不重新解释水平轨迹和姿态，而是在继承Stage2姿态、水平位置、水平速度和bias的基础上，仅对低频绝对高程进行校正。",
            "多成员场景下，先由各成员Stage2轨迹和可用RTKFIX样本构造统一距离域共享高程参考z_shared(s)，再将其映射到每个成员的Stage3-only优化中。",
            "v2.3 final策略使用5 mm高程包络、3 mm包络sigma、关闭center-pull，并加入Stage2垂向增量和jump形状继承，从而保持Stage3-Stage2差值主要为低频成分。"
        ],
    )

    add_heading(doc, "10. 历史验证与引用注意事项", 1)
    add_para(
        doc,
        "仓库文档中保留了历史lowfreq_delta_policy_default验证结果，用于说明IRI(stage2-stage3)的评价口径。"
        "这些数值是旧验证输出，不应直接等同于当前HEAD所有实验的最终指标。论文若要引用最终结果，应使用当前二进制、"
        "当前default_offline.cfg和当前Stage3 final策略在同一数据上重新运行，并保存输出目录。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录A. default_offline.cfg完整参数清单", 1)
    add_para(
        doc,
        f"本附录逐项列出{CFG_PATH.as_posix()}中解析到的全部默认参数，共{len(config_rows)}项。"
        "说明列只对关键参数给出具体解释，其余项按参数组归类。"
    )
    for group in GROUP_ORDER:
        rows = grouped.get(group, [])
        if not rows:
            continue
        add_heading(doc, group, 2)
        table_rows = [
            (
                key,
                fit_text(value, 72),
                bool_status(value),
                PARAMETER_NOTES.get(key, "完整默认配置项；结合本组说明和代码模块使用。"),
            )
            for key, value in rows
        ]
        add_table(doc, ["参数", "默认值", "状态", "说明"], table_rows, widths=[2.4, 1.6, 1.0, 2.5], size=6)

    add_heading(doc, "附录B. 主要源码依据", 1)
    source_rows = [
        ("CMakeLists.txt", "项目版本、库目标、三个可执行入口和测试目标。"),
        ("config/default_offline.cfg", "正式默认配置。"),
        ("include/offline_lc_minimal/common/Config.h", "OfflineRunnerConfig字段与代码兜底默认值。"),
        ("src/common/Config.cpp", "配置解析、枚举解析、单位转换、校验、ConfigToString。"),
        ("src/core/OfflineBatchRunner.cpp", "主因子图构建、stage递归调度、优化和诊断填充。"),
        ("src/core/OptimizationStagePolicy.cpp", "Stage1/Stage2子配置策略。"),
        ("src/core/Stage3HeightOptimizationPolicy.cpp", "Stage3 final高度优化策略覆盖。"),
        ("src/core/Stage1YawRefinementRunner.cpp", "Stage1 yaw精化流程。"),
        ("src/core/Stage2VelocityOptimizationRunner.cpp", "Stage2 source、body-y envelope、Stage2优化流程。"),
        ("src/core/Stage3VerticalReferenceOptimizationRunner.cpp", "内嵌Stage3 source/final调度。"),
        ("src/core/SharedVerticalReferenceBuilder.cpp", "共享距离域高程参考构造。"),
        ("src/stage3_runner/main.cpp", "Stage3-only共享参考入口。"),
    ]
    add_table(doc, ["文件", "使用依据"], source_rows, widths=[3.0, 4.5], size=8)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
